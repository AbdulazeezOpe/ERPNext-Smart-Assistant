import streamlit as st
from dotenv import load_dotenv
import os
from docx import Document
from gpt_parser import parse_doctype_prompt
from erp_api import create_erpnext_doctype, get_all_doctypes, get_records_for_doctype
from auth import check_login
from gpt_client import ask_gpt
from erp_api import create_workflow
from gpt_parser import parse_workflow_prompt
from erp_api import create_role, set_permission
from gpt_parser import parse_role_permission_prompt
from erp_api import create_notification
from gpt_parser import parse_notification_prompt
from erp_api import create_scheduled_reminder
from gpt_parser import parse_reminder_prompt
from erp_api import create_department
from gpt_parser import parse_department_prompt
from erp_api import create_user, assign_role_to_user
from gpt_parser import parse_user_assignment_prompt
from boq_parser import parse_boq_request
from erp_api import create_management_fee_rule, create_profit_sharing_rule
from gpt_parser import parse_financial_prompt
from erp_api import create_prf, get_pending_prfs
from gpt_parser import parse_prf_prompt
import re
from erp_api import add_vehicle, schedule_vehicle_maintenance, get_all_assets
from gpt_parser import parse_vehicle_prompt
from erp_api import generate_summary_report
from gpt_parser import parse_dashboard_prompt
from docx import Document







load_dotenv()

API_URL = os.getenv("ERP_BASE_URL")
API_KEY = os.getenv("ERP_API_KEY")
API_SECRET = os.getenv("ERP_API_SECRET")

st.set_page_config(page_title="ERPNext Smart Assistant", layout="wide")

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    st.title("🔐 ERPNext Assistant Login")
    password = st.text_input("Enter Password", type="password")
    if st.button("Login"):
        if check_login(password):
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("Incorrect password. Please try again.")

else:
    st.title("🤖 ERPNext Smart Assistant (Powered by GPT)")
    st.write("Welcome! This intelligent assistant allows you to: Create and manage Doctypes, Set up Workflows and Roles, Configure Notifications and Reminders, Analyze ERP data effortlessly and many more")

    action = st.radio("Choose an action", ["Ask GPT for Help", "Create Doctype Directly", "List and Explore Available Doctypes", "📚 Documentation"])

    if action == "Ask GPT for Help":
        gpt_prompt = st.text_area("Ask a question or request guidance:", height=150)

        if st.button("Ask GPT"):
            if gpt_prompt:
                with st.spinner("GPT is thinking..."):
                    gpt_response = ask_gpt(gpt_prompt)
                    st.success(gpt_response)

                    # Check if prompt mentions notification setup
                    if "notify" in gpt_prompt.lower() or "send email" in gpt_prompt.lower():
                        st.info("Detected notification setup request. Parsing details...")

                        from gpt_parser import parse_notification_prompt
                        subject, document_type, condition, message = parse_notification_prompt(gpt_prompt)

                        auth = (API_KEY, API_SECRET)

                        notification_result = create_notification(API_URL, auth, subject, document_type, condition, message)

                        if "data" in notification_result:
                            st.success(f"✅ Notification for **{document_type} approval** setup successfully! 📩")
                        else:
                            st.error("❌ Failed to create Notification. Check prompt or ERP access.")

                    
                    # Detect Inventory & Supplier Related Actions
                    if any(word in gpt_prompt.lower() for word in ["inventory", "item", "supplier", "stock"]):
                        st.info("Detected inventory/supplier management request. Parsing details...")

                        from gpt_parser import parse_inventory_prompt
                        from erp_api import add_inventory_item, get_low_stock_items, add_supplier

                        item_name, quantity, department_name, supplier_name, contact_number = parse_inventory_prompt(gpt_prompt)

                        auth = (API_KEY, API_SECRET)

                        if "item" in gpt_prompt.lower():
                            if item_name and quantity:
                                if st.checkbox(f"✅ Confirm adding Inventory Item: {item_name} (Qty: {quantity})?"):
                                    result = add_inventory_item(API_URL, auth, item_name, quantity, department_name)
                                    if "data" in result:
                                        st.success(f"📦 Inventory item '{item_name}' added successfully!")
                                    else:
                                        st.error("❌ Failed to add inventory item. Check input details.")
                                else:
                                    st.warning("⚡ Please confirm before adding inventory item.")

                        elif "supplier" in gpt_prompt.lower():
                            if supplier_name:
                                if st.checkbox(f"✅ Confirm adding Supplier: {supplier_name}?"):
                                    result = add_supplier(API_URL, auth, supplier_name, contact_number)
                                    if "data" in result:
                                        st.success(f"🏢 Supplier '{supplier_name}' added successfully!")
                                    else:
                                        st.error("❌ Failed to add supplier. Check input details.")
                                else:
                                    st.warning("⚡ Please confirm before adding supplier.")

                        elif "low stock" in gpt_prompt.lower():
                            low_stock = get_low_stock_items(API_URL, auth)
                            if low_stock:
                                st.write("### Low Stock Items")
                                st.dataframe(low_stock)
                            else:
                                st.warning("No low stock items found.")

                    # Detect Dashboard/Reports Related Actions
                    if "dashboard" in gpt_prompt.lower() or "report" in gpt_prompt.lower() or "summary" in gpt_prompt.lower():
                        st.info("Detected report or dashboard request. Parsing details...")

                        from gpt_parser import parse_dashboard_prompt
                        from erp_api import generate_summary_report

                        target_audience, report_type, frequency = parse_dashboard_prompt(gpt_prompt)

                        auth = (API_KEY, API_SECRET)

                        if st.button("📄 Generate Report Now"):
                            with st.spinner("Generating report..."):
                                summary = generate_summary_report(API_URL, auth)

                                # Display the summary
                                st.write("### 📋 Report Summary")
                                st.json(summary)

                                # Optionally download as Word document
                                from docx import Document

                                doc = Document()
                                doc.add_heading('ERPNext Summary Report', 0)

                                if target_audience:
                                    doc.add_paragraph(f"Target Audience: {target_audience}")
                                if frequency:
                                    doc.add_paragraph(f"Frequency: {frequency}")

                                doc.add_paragraph("Summary:")

                                for key, value in summary.items():
                                    doc.add_paragraph(f"{key}: {value}")

                                report_filename = "erp_summary_report.docx"
                                doc.save(report_filename)

                                with open(report_filename, "rb") as file:
                                    st.download_button(
                                        label="📥 Download Summary Report",
                                        data=file,
                                        file_name=report_filename,
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                    )

                    
                    
                    # Detect Claim Creation or Query
                    if "claim" in gpt_prompt.lower():
                        st.info("Detected claim request. Parsing details...")

                        from gpt_parser import parse_claim_prompt
                        from erp_api import create_claim, get_claims

                        project_name, claim_name, amount = parse_claim_prompt(gpt_prompt)

                        auth = (API_KEY, API_SECRET)

                        if "submit" in gpt_prompt.lower() or "create" in gpt_prompt.lower():
                            # User wants to create a Claim
                            if project_name and claim_name and amount:
                                result = create_claim(API_URL, auth, project_name, claim_name, amount)
                                if "data" in result:
                                    st.success(f"📄 Claim '{claim_name}' for project '{project_name}' created successfully (RM {amount})!")
                                else:
                                    st.error("❌ Failed to create claim. Check input details.")
                            else:
                                st.error("❌ Could not detect project, claim name, or amount properly.")

                        elif "due" in gpt_prompt.lower() or "status" in gpt_prompt.lower() or "list" in gpt_prompt.lower():
                            # User wants to check claims
                            claims = get_claims(API_URL, auth, project_name)

                            if claims:
                                st.write(f"### Claims for {project_name}")
                                st.dataframe(claims)
                            else:
                                st.warning(f"No claims found for project {project_name}.")

                                        
                    
                    
                    # Detect Financial Automation (Management Fee or Profit Sharing)
                    if ("management fee" in gpt_prompt.lower()) or ("profit sharing" in gpt_prompt.lower()) or ("profit rule" in gpt_prompt.lower()):
                        st.info("Detected financial automation request. Parsing details...")

                        from gpt_parser import parse_financial_prompt
                        from erp_api import create_management_fee_rule, create_profit_sharing_rule

                        action, department, role, amount = parse_financial_prompt(gpt_prompt)

                        auth = (API_KEY, API_SECRET)

                        if action == "management_fee":
                            if department and amount:
                                result = create_management_fee_rule(API_URL, auth, department, amount)
                                if "data" in result:
                                    st.success(f"💰 Management fee of RM {amount} set for {department} department!")
                                else:
                                    st.error("❌ Failed to create management fee rule. Check details.")
                            else:
                                st.error("❌ Could not detect department or amount for management fee setup.")

                        elif action == "profit_sharing":
                            if role and amount:
                                # Assume department if missing (optional to improve later)
                                department = department or "General"
                                result = create_profit_sharing_rule(API_URL, auth, department, role, amount)
                                if "data" in result:
                                    st.success(f"🎯 Profit sharing rule: {amount}% to {role} set successfully!")
                                else:
                                    st.error("❌ Failed to create profit sharing rule. Check details.")
                            else:
                                st.error("❌ Could not detect role or percentage for profit sharing rule.")

                                        
                    # Detect PRF Creation
                    if "prf" in gpt_prompt.lower() and ("create" in gpt_prompt.lower() or "submit" in gpt_prompt.lower()):
                        st.info("Detected PRF creation request. Parsing details...")

                        from gpt_parser import parse_prf_prompt
                        from erp_api import create_prf

                        project_name, item_name, quantity = parse_prf_prompt(gpt_prompt)

                        auth = (API_KEY, API_SECRET)

                        if project_name and item_name and quantity:
                            result = create_prf(API_URL, auth, project_name, item_name, quantity)
                            if "data" in result:
                                st.success(f"📄 PRF created for {item_name} under {project_name} (Qty: {quantity}) successfully!")
                            else:
                                st.error("❌ Failed to create PRF. Check details.")
                        else:
                            st.error("❌ Could not detect project, item or quantity properly.")
                   
                    
                    # Detect Pending PRF Listing
                    if "list" in gpt_prompt.lower() and "pending prf" in gpt_prompt.lower():
                        st.info("Detected request to list pending PRFs.")

                        from erp_api import get_pending_prfs

                        auth = (API_KEY, API_SECRET)

                        # Try to detect department from text
                        dept_match = re.search(r"for\s+([\w\s]+?)\s+department", gpt_prompt.lower())
                        department_name = dept_match.group(1).strip().title() if dept_match else None

                        pending_prfs = get_pending_prfs(API_URL, auth, department_name)

                        if pending_prfs:
                            st.write(f"### Pending PRFs{' for ' + department_name if department_name else ''}")
                            st.dataframe(pending_prfs)
                        else:
                            st.warning("No pending PRFs found or unable to fetch records.")

                    
                    
                    # Check if prompt mentions adding a user
                    if "user" in gpt_prompt.lower() and ("add" in gpt_prompt.lower() or "create" in gpt_prompt.lower()):
                        st.info("Detected user creation and role assignment request. Parsing details...")

                        from gpt_parser import parse_user_assignment_prompt
                        user_name, department_name, role_name = parse_user_assignment_prompt(gpt_prompt)

                        auth = (API_KEY, API_SECRET)

                        # Assume email based on name for demo (you can improve this later)
                        email = f"{user_name.lower()}@example.com"

                        # Step 1: Create the User
                        user_result = create_user(API_URL, auth, user_name, email)

                        if "data" in user_result:
                            st.success(f"✅ User '{user_name}' created successfully!")

                            # Step 2: Assign Role
                            role_result = assign_role_to_user(API_URL, auth, email, role_name)

                            if "data" in role_result:
                                st.success(f"🎯 Role '{role_name}' assigned to '{user_name}' successfully!")
                            else:
                                st.error("❌ Failed to assign role. Check role name or permissions.")

                        else:
                            st.error("❌ Failed to create user. User may already exist or check API access.")

                    
                    # Check if prompt mentions BOQ creation
                    if "boq" in gpt_prompt.lower() and ("create" in gpt_prompt.lower() or "add" in gpt_prompt.lower()):
                        st.info("Detected BOQ creation request. Parsing details...")

                        from gpt_parser import parse_boq_creation_prompt
                        budget_amount_rm, item_name, quantity, price = parse_boq_creation_prompt(gpt_prompt)

                        print("👉 DEBUG: Project Name parsed:", project_name)
                        print("👉 DEBUG: Item Name parsed:", item_name)
                        print("👉 DEBUG: Quantity parsed:", quantity)
                        print("👉 DEBUG: Price parsed:", price)

                        auth = (API_KEY, API_SECRET)

                        from erp_api import create_boq_entry

                        # ✅ If we reach here, project exists and is valid → Create BOQ
                        boq_result = create_boq_entry(API_URL, auth, budget_amount_rm, item_name, quantity, price)
                        print("🧠 Project API Response:", boq_result)

                        if "data" in boq_result:
                            boq_data = boq_result["data"]
                            st.success(
                                f"📦 BOQ item **'{boq_data.get('boq_item_description')}'** created successfully '**!"
                            )
                        else:
                            st.error("❌ Failed to create BOQ item. Check prompt or ERP access.")

                                
                               


                           
                    
                    
                    # Check if prompt mentions department creation
                    if "department" in gpt_prompt.lower() and ("create" in gpt_prompt.lower() or "add" in gpt_prompt.lower()):
                        st.info("Detected department creation request. Parsing details...")

                        from gpt_parser import parse_department_prompt
                        
                        department_name, parent_department = parse_department_prompt(gpt_prompt)
                        print("👉 DEBUG: Department Name parsed:", department_name)
                        print("👉 DEBUG: Parent Department parsed:", parent_department)

                        auth = (API_KEY, API_SECRET)
                        company_name = "S&I Urban Designers"  # <-- Your company

                        from erp_api import create_department
                        department_result = create_department(API_URL, API_KEY, API_SECRET, department_name, company_name, parent_department)

                        if "data" in department_result:
                            department_data = department_result["data"]
                            st.success(
                                f"🏢 Department **'{department_data.get('department_name')}'** created successfully under **'{department_data.get('parent_department')}'**!"
                            )
                        else:
                            st.error("❌ Failed to create department. Check prompt or ERP access.")



                    
                    # Detect Project Creation or Assignment
                    if "project" in gpt_prompt.lower() and ("create" in gpt_prompt.lower() or "assign" in gpt_prompt.lower()):
                        st.info("Detected project management request. Parsing details...")

                        from gpt_parser import parse_project_prompt
                        from erp_api import create_project, assign_project_roles

                        project_name, expected_end_date, estimated_costing, assignments = parse_project_prompt(gpt_prompt)
                        auth = (API_KEY, API_SECRET)

                        # Show parsed project data
                        if project_name:
                            st.info(f"🧾 Parsed Project:\n• Name: {project_name}\n• End Date: {expected_end_date or 'N/A'}\n• Cost: {estimated_costing or 'N/A'}")

                        # Handle project creation
                        if "create" in gpt_prompt.lower():
                            if not project_name:
                                st.error("❌ Project name is missing. Please rephrase your instruction.")
                            else: 
                                result = create_project(API_URL, auth, project_name, expected_end_date, estimated_costing)
                                print("🧠 Project API Response:", result)

                                if result and "data" in result:
                                    st.success(f"🏗️ Project '{project_name}' created successfully!")
                                else:
                                    st.error("❌ Failed to create project. Check input details or ERP connection.")

                        



                    
                    # Detect Vehicle/Asset Related Actions
                    if "vehicle" in gpt_prompt.lower() or "asset" in gpt_prompt.lower():
                        st.info("Detected vehicle/asset management request. Parsing details...")

                        from gpt_parser import parse_vehicle_prompt
                        from erp_api import add_vehicle, schedule_vehicle_maintenance, get_all_assets

                        vehicle_name, department_name, maintenance_interval_months = parse_vehicle_prompt(gpt_prompt)

                        auth = (API_KEY, API_SECRET)

                        if "add" in gpt_prompt.lower() or "new vehicle" in gpt_prompt.lower():
                            if vehicle_name:
                                if st.checkbox(f"✅ Confirm adding Vehicle: {vehicle_name}?"):
                                    result = add_vehicle(API_URL, auth, vehicle_name, department_name)
                                    if "data" in result:
                                        st.success(f"🚗 Vehicle '{vehicle_name}' added successfully!")
                                    else:
                                        st.error("❌ Failed to add vehicle. Check input details.")
                                else:
                                    st.warning("⚡ Please confirm before adding vehicle.")

                        elif "schedule" in gpt_prompt.lower() and "maintenance" in gpt_prompt.lower():
                            if vehicle_name and maintenance_interval_months:
                                if st.checkbox(f"✅ Confirm scheduling maintenance every {maintenance_interval_months} months for {vehicle_name}?"):
                                    result = schedule_vehicle_maintenance(API_URL, auth, vehicle_name, maintenance_interval_months)
                                    if "data" in result:
                                        st.success(f"🛠️ Maintenance scheduled for '{vehicle_name}' every {maintenance_interval_months} months!")
                                    else:
                                        st.error("❌ Failed to schedule maintenance.")
                                else:
                                    st.warning("⚡ Please confirm before scheduling maintenance.")

                        elif "list" in gpt_prompt.lower() and ("vehicles" in gpt_prompt.lower() or "assets" in gpt_prompt.lower()):
                            assets = get_all_assets(API_URL, auth)
                            if assets:
                                st.write("### Company Vehicles/Assets")
                                st.dataframe(assets)
                            else:
                                st.warning("No vehicles/assets found.")

                    
                    
                    # Detect HR Related Actions (Contract, Salary Advance, Leave)
                    if any(word in gpt_prompt.lower() for word in ["contract", "salary advance", "leave balance", "track leave"]):
                        st.info("Detected HR-related request. Parsing details...")

                        from gpt_parser import parse_hr_prompt
                        from erp_api import generate_contract, get_salary_advances, get_leave_balance

                        employee_name, monthly_salary, start_date = parse_hr_prompt(gpt_prompt)

                        auth = (API_KEY, API_SECRET)

                        if "contract" in gpt_prompt.lower():
                            if employee_name and monthly_salary and start_date:
                                if st.checkbox(f"✅ Confirm contract generation for {employee_name}?"):
                                    result = generate_contract(API_URL, auth, employee_name, monthly_salary, start_date)
                                    if "data" in result:
                                        st.success(f"📝 Employment contract created for {employee_name} (RM {monthly_salary}/month) starting {start_date}!")
                                    else:
                                        st.error("❌ Failed to create contract. Check input details.")
                                else:
                                    st.warning("⚡ Please confirm before generating contract.")

                        elif "salary advance" in gpt_prompt.lower():
                            if employee_name:
                                advances = get_salary_advances(API_URL, auth, employee_name)
                                if advances:
                                    st.write(f"### Salary Advances for {employee_name}")
                                    st.dataframe(advances)
                                else:
                                    st.warning(f"No salary advance records found for {employee_name}.")
                            else:
                                advances = get_salary_advances(API_URL, auth)
                                if advances:
                                    st.write("### All Salary Advances")
                                    st.dataframe(advances)
                                else:
                                    st.warning("No salary advance records found.")

                        elif "leave" in gpt_prompt.lower():
                            if employee_name:
                                leaves = get_leave_balance(API_URL, auth, employee_name)
                                if leaves:
                                    st.write(f"### Leave Balance for {employee_name}")
                                    st.dataframe(leaves)
                                else:
                                    st.warning(f"No leave balance records found for {employee_name}.")
                            else:
                                st.warning("Please specify an employee name for leave balance checking.")

                    
                    
                    # Check if prompt mentions role creation
                    if "create" in gpt_prompt.lower() and "role" in gpt_prompt.lower():
                        st.info("Detected role creation request. Parsing details...")

                        from gpt_parser import parse_role_permission_prompt
                        role_name, allowed_doctypes, restricted_doctypes = parse_role_permission_prompt(gpt_prompt)

                        auth = (API_KEY, API_SECRET)

                        # Create the Role
                        role_result = create_role(API_URL, auth, role_name)
                        if "data" in role_result:
                            st.success(f"✅ Role '{role_name}' created successfully!")

                            # Set Permissions for Allowed Doctypes
                            for doctype in allowed_doctypes:
                                perm_result = set_permission(API_URL, auth, doctype, role_name, read=1, write=1, create=1)
                                if "data" in perm_result:
                                    st.success(f"✅ Role **{role_name}** created and permissions assigned! 🎯")

                            # Optionally: You can add logic later to restrict permissions if needed
                            if restricted_doctypes:
                                st.warning(f"ℹ️ Restricted Doctypes detected: {', '.join(restricted_doctypes)}. Please manually restrict via Permission Manager (recommended for security).")

                        else:
                            st.error("❌ Failed to create Role. Check prompt or ERP access.")

                    
                    
                    
                    # Check if prompt mentions workflow creation
                    if "approval workflow" in gpt_prompt.lower() or "create workflow" in gpt_prompt.lower():
                        st.info("Detected workflow creation request. Parsing details...")

                        # Parse workflow
                        from gpt_parser import parse_workflow_prompt
                        workflow_name, document_type, states, transitions = parse_workflow_prompt(gpt_prompt)

                        # Create workflow
                        auth = (API_KEY, API_SECRET)
                        result = create_workflow(API_URL, auth, workflow_name, document_type, states, transitions)

                        if "data" in result:
                            st.success(f"✅ Successfully created Workflow: **{workflow_name}** for {document_type}. 🚀")
                        else:
                            st.error("❌ Failed to create workflow. Check prompt or ERP access.")

            else:
                st.error("Please enter a question or instruction.")

    elif action == "Setup Reminder":
        reminder_prompt = st.text_area("Describe the reminder you want to set (e.g., 'Send weekly reminders for all pending PRFs')", height=150)

        if st.button("Setup Reminder"):
            if reminder_prompt:
                with st.spinner("Parsing and setting up scheduled reminder..."):
                    document_type, condition, message, cron = parse_reminder_prompt(reminder_prompt)

                    auth = (API_KEY, API_SECRET)
                    result = create_scheduled_reminder(API_URL, auth, document_type, condition, message, cron)

                    if "data" in result:
                        st.success(f"✅ Scheduled Reminder for **{document_type}** set to run automatically! ⏰")
                    else:
                        st.error("❌ Failed to set up reminder. Please check your prompt or ERP access.")
            else:
                st.error("Please describe the reminder you want to set.")

    
    
    elif action == "Create Doctype Directly":
        prompt = st.text_area("Enter a prompt like: 'Create a new BOQ Doctype with fields...'")

        if st.button("Create Doctype"):
            if not prompt:
                st.error("Please enter a prompt.")
            else:
                with st.spinner("Parsing and creating Doctype..."):
                    doctype_name, fields = parse_doctype_prompt(prompt)
                    st.write("🛠 Parsed Doctype Name:", doctype_name)
                    st.write("🛠 Parsed Fields:", fields)
                    auth = (API_KEY, API_SECRET)
                    result = create_erpnext_doctype(API_URL, auth, doctype_name, fields)
                    if "data" in result:
                        st.success(f"✅ Doctype '{doctype_name}' created successfully!")
                    else:
                        st.error("❌ Failed to create Doctype. Check the prompt or API access.")

    elif action == "List and Explore Available Doctypes":
        with st.spinner("Fetching available Doctypes..."):
            auth = (API_KEY, API_SECRET)
            doctypes = get_all_doctypes(API_URL, auth)
            if doctypes:
                selected_doctype = st.selectbox("Select a Doctype to view records:", doctypes)

                if selected_doctype:
                    with st.spinner(f"Fetching records for {selected_doctype}..."):
                        records = get_records_for_doctype(API_URL, auth, selected_doctype)
                        if records:
                            st.write(f"### Records for {selected_doctype}")
                            st.dataframe(records)

                            
                            if selected_doctype.lower() in ["project boq", "boq breakdown"]:
                                boq_question = st.text_input("Optional: Ask about BOQ usage (e.g., 'Show balance for Cement in Project ABC')")

                                if boq_question:
                                    action, project_name, boq_item = parse_boq_request(boq_question)

                                    if action == "show_balance":
                                        if project_name and boq_item:
                                            matching = [r for r in records if r.get("project_name", "").lower() == project_name.lower() and boq_item.lower() in r.get("boq_item_description", "").lower()]
                                            if matching:
                                                st.success(f"Balance for {boq_item} in {project_name}: {matching[0].get('balance_amount', 'N/A')}")
                                            else:
                                                st.warning(f"No matching BOQ item found for {boq_item} in {project_name}.")
                                        else:
                                            st.warning("Could not detect project name or BOQ item correctly.")

                                    elif action == "list_boq_items":
                                        if project_name:
                                            project_boqs = [r for r in records if r.get("project_name", "").lower() == project_name.lower()]
                                            if project_boqs:
                                                st.write(f"BOQ Items for {project_name}:")
                                                for boq in project_boqs:
                                                    st.write(f"- {boq.get('boq_item_description', 'Unnamed Item')}")
                                            else:
                                                st.warning(f"No BOQ items found for {project_name}.")
                                        else:
                                            st.warning("Could not detect project name correctly.")

                            
                            
                            
                            custom_question = st.text_area("Optional: Ask GPT a specific question about these records", placeholder="e.g., Find all requests above $10,000")

                            if st.button("Analyze Records with GPT"):
                                record_text = "\n".join([str(record) for record in records])
                                if custom_question.strip():
                                    query_prompt = f"{custom_question}\n\nHere are the records:\n{record_text}"
                                else:
                                    query_prompt = f"Analyze the following ERPNext {selected_doctype} records and summarize any important insights or risks you detect:\n\n{record_text}"

                                with st.spinner("GPT is analyzing the records..."):
                                    analysis = ask_gpt(query_prompt)
                                    st.success(analysis)

                                    # Download option
                                    doc = Document()
                                    doc.add_heading("GPT Analysis Report", 0)
                                    doc.add_paragraph(analysis)
                                    report_filename = "gpt_analysis_report.docx"
                                    doc.save(report_filename)

                                    with open(report_filename, "rb") as file:
                                        st.download_button(
                                            label="📥 Download GPT Report",
                                            data=file,
                                            file_name=report_filename,
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                        )
                        else:
                            st.warning("No records found or access restricted.")
            else:
                st.warning("Couldn't fetch Doctypes or no access.")

    elif action == "📚 Documentation":
        st.title("📚 ERPNext Smart Assistant - Full User Guide")

        st.header("🔹 Welcome")
        st.write("""
        Welcome to the ERPNext Smart Assistant powered by GPT and Python! 
        This system automates your ERP workflows and allows you to manage operations by simply typing instructions in natural language.
        """)

        st.header("🔐 How to Login")
        st.write("""
        1. Open the app.
        2. Enter your password (assigned by Admin).
        3. Once authenticated, you will access the Smart Assistant interface.
        """)

        st.header("🧠 Ask GPT for Help")
        st.write("""
        This is your smart communication zone where you can type instructions naturally.

        **How to Use:**
        - Type a clear instruction into the box.
        - Example prompts:
        - "Create a new BOQ Doctype with fields Project Name, Quantity, Unit Price."
        - "List pending PRFs for Finance department."
        - "Generate monthly financial report for Director."
        - Click the button to send your instruction.
        - GPT will understand, process, and execute your request live!

        **Best Tip:**  
        Be as clear as possible when typing instructions for best results!
        """)

        st.header("🏗 Create Doctype Directly")
        st.write("""
        This allows you to directly create new ERP structures without needing to manually enter ERPNext backend.

        **How to Use:**
        1. Go to "Create Doctype Directly".
        2. Enter a prompt like:
        - "Create a new Doctype called 'Subcontractor Performance' with fields Name, Rating, Project."
        3. Click "Create Doctype".
        4. The system automatically builds it in ERPNext.

        **Best Tip:**  
        Always mention both the Doctype Name and its fields clearly.
        """)

        st.header("📂 List and Explore Available Doctypes")
        st.write("""
        This lets you view all existing ERPNext modules and records.

        **How to Use:**
        1. Select "List and Explore Available Doctypes".
        2. The app fetches all available Doctypes from ERPNext.
        3. Select a Doctype from the dropdown (e.g., PRF, Claim, Project).
        4. View and analyze records inside.
        5. Optionally, ask GPT questions about the records (e.g., "Highlight all PRFs over RM 10,000").

        **Best Tip:**  
        If records are not showing, ensure that ERP permissions are correctly set for your API user.
        """)

        st.header("🏢 Department and Role Setup")
        st.write("""
        Manage departments and assign users automatically.

        **How to Use:**
        1. In "Ask GPT for Help", type:
        - "Create department called Facility Maintenance."
        - "Add user Ali to Facility Maintenance as HOD."
        2. Confirm the action through checkboxes.
        3. New departments or user assignments will be created instantly.

        **Best Tip:**  
        Mention the department name exactly as you want it to appear in ERPNext.
        """)

        st.header("📑 PRF Management (Purchase Request Form)")
        st.write("""
        Create, approve, and track PRFs easily.

        **How to Use:**
        1. Ask:
        - "Create a PRF for Project ABC, Item: Paint, Quantity: 100."
        2. GPT will create the PRF automatically.
        3. You can also list pending PRFs:
        - "List all pending PRFs for Finance."

        **Best Tip:**  
        Always specify Project Name, Item Name, and Quantity when creating PRFs.
        """)

        st.header("📃 Claims Management")
        st.write("""
        Submit and track project claims.

        **How to Use:**
        1. Ask:
        - "Submit claim for Project XYZ, Amount RM 50,000."
        2. To list claims:
        - "List all approved claims."

        **Best Tip:**  
        Ensure Project Name and Claim Amount are clear inside the prompt.
        """)

        st.header("🏢 Project Management")
        st.write("""
        Manage projects including creation and team assignment.

        **How to Use:**
        1. Ask:
        - "Create project ABC, Start May 1st, Budget RM 300,000."
        2. Assign managers:
        - "Assign HOD and Project Manager to Project ABC."

        **Best Tip:**  
        Mention start date and budget if possible during project creation.
        """)

        st.header("💰 Financial Management (Management Fees and Profit Sharing)")
        st.write("""
        Automate department management fees and profit distribution.

        **How to Use:**
        1. Ask:
        - "Apply RM 10,000 management fee monthly to all departments."
        - "Set profit sharing rule 10% to HOD after project completion."

        **Best Tip:**  
        Clearly state the percentage and conditions if configuring profit sharing.
        """)

        st.header("🚗 Vehicle and Asset Management")
        st.write("""
        Manage vehicles, assign them to departments, and automate maintenance scheduling.

        **How to Use:**
        1. Add vehicles:
        - "Add vehicle Toyota Hilux to Facility Maintenance."
        2. Schedule maintenance:
        - "Schedule maintenance for Toyota Hilux every 3 months."

        **Best Tip:**  
        Always mention Vehicle Name and Department clearly.
        """)

        st.header("📦 Inventory & Supplier Management")
        st.write("""
        Add stock items, list low stock, and manage suppliers.

        **How to Use:**
        1. Add inventory item:
        - "Add inventory item Paint, Quantity 100, Department Signage."
        2. Add supplier:
        - "Add supplier ABC Materials, Contact 0123456789."
        3. List low stock items:
        - "List all low stock items."

        **Best Tip:**  
        Quantity and Department are key fields for inventory items.
        """)

        st.header("📚 Dashboard & Reports Generation")
        st.write("""
        Create summary reports and dashboards automatically.

        **How to Use:**
        1. Ask:
        - "Generate monthly financial report and send to Director."
        2. The system will generate live summary reports based on ERP data.
        3. Download reports as Word documents instantly.

        **Best Tip:**  
        Mention target audience (HOD, Director) and frequency (monthly, weekly) inside your prompt.
        """)

        st.header("🔔 Notification System")
        st.write("""
        System will automatically send reminders for:
        - Pending PRFs
        - Pending Claims
        - Project deadlines

        No manual setup needed once workflows are active!
        """)


        st.header("⚡ Tips for Best Use")
        st.write("""
        - Be as clear and structured as possible when typing instructions.
        - Always double-check field names like Project Name, Department Name, etc.
        - Use the in-app confirmation checkboxes before performing major actions.
        - Start small, then scale up the automation step-by-step.
        """)

        st.success("🎯 You are now fully equipped to master the ERPNext Smart Assistant!")
